import configparser
import json
import logging
import os
import shutil
import tempfile
import threading
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from typing import Dict, List, Optional, Tuple

import pyperclip
import textract
from docx import Document

from component.label_frame_bottom_left import LabelFrameBottomLeft
from component.label_frame_right import LabelFrameRight
from component.label_frame_top_left import LabelFrameTopLeft
from config import settings


class PseudonymStore:
    """
    Класс для работы со словарём псевдонимов.

    Хранит:
    - путь к JSON-файлу словаря;
    - пары «оригинал → псевдоним» и «псевдоним → оригинал».

    Предоставляет методы для:
    - загрузки и сохранения словаря из/в файл;
    - генерации новых псевдонимов;
    - формирования упорядоченных списков замен для анонимизации и деанонимизации.
    """

    def __init__(self, path: Path) -> None:
        self.path = path
        self.anon_dict: Dict[str, str] = {}
        self.deanon_dict: Dict[str, str] = {}

    def load(self) -> None:
        pseudos_file = self.path
        try:
            if pseudos_file.exists():
                with open(pseudos_file, "r", encoding="utf-8") as f:
                    self.anon_dict = json.load(f)
                self.deanon_dict = {v: k for k, v in self.anon_dict.items()}
                logging.info(f"Загружено {len(self.anon_dict)} псевдонимов")
        except Exception as e:
            logging.error(f"Ошибка загрузки pseudos.json: {e}")
            messagebox.showerror(
                "Ошибка", f"Не удалось загрузить словарь псевдонимов: {e}"
            )

    def save(self) -> None:
        pseudos_file = self.path
        backup_file = pseudos_file.with_suffix(".json.bak")

        try:
            if pseudos_file.exists():
                shutil.copy2(pseudos_file, backup_file)

            with tempfile.NamedTemporaryFile(
                mode="w", encoding="utf-8", dir=pseudos_file.parent, delete=False
            ) as tmp:
                json.dump(self.anon_dict, tmp, ensure_ascii=False, indent=2)
                tmp_path = tmp.name

            os.replace(tmp_path, pseudos_file)
            logging.info("Словарь псевдонимов сохранён")
        except Exception as e:
            logging.error(f"Ошибка сохранения pseudos.json: {e}")
            messagebox.showerror(
                "Ошибка", f"Не удалось сохранить словарь псевдонимов: {e}"
            )

    def generate_pseudonym(self, initial_pseudo: Optional[str] = None) -> str:
        max_num = 0
        initial_pseudo = initial_pseudo or "ЗАМЕНА"
        for pseudo in self.deanon_dict.keys():
            if pseudo.startswith(initial_pseudo):
                try:
                    num_str = pseudo[len(initial_pseudo) :]
                    num = int(num_str)
                    max_num = max(max_num, num)
                except ValueError:
                    continue
        new_num = max_num + 1
        return f"{initial_pseudo}{new_num:03d}"

    def get_replacement_order(self) -> List[Tuple[str, str]]:
        items = list(self.anon_dict.items())
        items.sort(key=lambda x: (-len(x[0]), x[0]))
        return items

    def get_deanonymization_order(self) -> List[Tuple[str, str]]:
        items = [
            (f"[{pseudo}]", original) for pseudo, original in self.deanon_dict.items()
        ]
        items.sort(key=lambda x: (-len(x[0]), x[0]))
        return items


class FileProcessor:
    """
    Класс для работы с файлами (.txt, .docx, .doc).

    Отвечает за:
    - чтение исходного файла;
    - применение списка замен;
    - сохранение результата и генерацию имени выходного файла.
    """

    def __init__(self) -> None:
        # Пока без состояния, оставляем на будущее.
        pass

    def _process_docx_with_mapping_and_save(
        self,
        file_path: Path,
        mapping: List[Tuple[str, str]],
        suffix: str,
    ) -> str:
        def apply_mapping_to_text(text: str, mapping: List[Tuple[str, str]]) -> str:
            # mapping уже должен быть отсортирован снаружи
            for src, dst in mapping:
                if src:
                    text = text.replace(src, dst)
            return text

        def replace_in_paragraph(paragraph, mapping: List[Tuple[str, str]]):
            old = paragraph.text
            new = apply_mapping_to_text(old, mapping)
            if new != old:
                # убиваем старые run'ы, чтобы не ломать форматирование сильнее, чем надо
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = new
                else:
                    paragraph.add_run(new)

        def replace_in_table(table, mapping: List[Tuple[str, str]]):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p, mapping)

        try:
            doc = Document(str(file_path))
        except Exception as e:
            user_msg = (
                "Не удалось открыть файл как документ Word.\n\n"
                "Возможные причины:\n"
                "• файл открыт и ещё не полностью сохранён в другой программе;\n"
                "• файл удалён, переименован или перемещён;\n"
                "• файл повреждён и не читается как .docx.\n\n"
                "Проверьте, что файл существует, закрыт во всех программах "
                "и попробуйте ещё раз."
            )
            raise ValueError(user_msg)

        # Основной текст
        for p in doc.paragraphs:
            replace_in_paragraph(p, mapping)
        # Таблицы
        for table in doc.tables:
            replace_in_table(table, mapping)
        # Колонтитулы всех секций
        for section in doc.sections:
            for p in section.header.paragraphs:
                replace_in_paragraph(p, mapping)
            for p in section.footer.paragraphs:
                replace_in_paragraph(p, mapping)
            for table in section.header.tables:
                replace_in_table(table, mapping)
            for table in section.footer.tables:
                replace_in_table(table, mapping)

        output_path = self.get_output_path(file_path, suffix, ".docx")
        try:
            doc.save(str(output_path))
        except Exception as e:
            raise ValueError(f"Ошибка сохранения .docx файла: {e}")

        return f"Файл сохранён как: {output_path.name}"

    def process_doc_file(
        self,
        file_path: Path,
        replacements: List[Tuple[str, str]],
        anonymize: bool = True,
    ) -> str:
        if textract is None:
            raise ValueError(
                "Поддержка .doc требует textract; "
                "конвертируйте в .docx или установите textract"
            )

        try:
            text = textract.process(str(file_path)).decode("utf-8")
        except Exception as e:
            raise ValueError(f"Ошибка чтения .doc файла: {e}")

        if anonymize:
            for original, pseudo in replacements:
                text = text.replace(original, f"[{pseudo}]")
        else:
            for pseudo_pattern, original in replacements:
                text = text.replace(pseudo_pattern, original)

        suffix = "_аноним" if anonymize else "_деаном"
        output_path = self.get_output_path(file_path, suffix, ".txt")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

        return f"Файл сохранён как: {output_path.name}"

    def get_output_path(
        self, original_path: Path, suffix: str, force_ext: Optional[str] = None
    ) -> Path:
        ext = force_ext if force_ext else original_path.suffix
        base_name = original_path.stem

        counter = 0
        while True:
            if counter == 0:
                new_name = f"{base_name}{suffix}{ext}"
            else:
                new_name = f"{base_name}{suffix}{counter:03d}{ext}"

            output_path = original_path.parent / new_name
            if not output_path.exists():
                return output_path
            counter += 1

    def process_txt_file(
        self, file_path: Path, replacements: List[Tuple[str, str]]
    ) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="cp1251") as f:
                    content = f.read()
            except UnicodeDecodeError as e:
                raise ValueError(f"Не удалось прочитать файл: {e}")

        for original, pseudo in replacements:
            content = content.replace(original, f"[{pseudo}]")

        output_path = self.get_output_path(file_path, "_аноним")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

        return f"Файл сохранён как: {output_path.name}"

    def process_docx_file(
        self,
        file_path: Path,
        replacements: List[Tuple[str, str]],
        anonymize: bool = True,
    ) -> str:
        # anonymize = True  -> анонимизируем и сохраняем .docx
        # anonymize = False -> деанонимизируем и возвращаем текст

        if anonymize:
            # пары (оригинал, псевдоним) -> (оригинал, [псевдоним])
            mapping = [(orig, f"[{pseudo}]") for orig, pseudo in replacements]
            return self._process_docx_with_mapping_and_save(
                file_path=file_path,
                mapping=mapping,
                suffix="_аноним",
            )

        # -------- режим деанонимизации: вернуть текст в буфер --------
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        def apply_mapping_to_text(text: str, mapping: List[Tuple[str, str]]) -> str:
            for src, dst in mapping:
                if src:
                    text = text.replace(src, dst)
            return text

        def iter_block_items(container):
            """
            Идём по детям контейнера (документ, ячейка, header/footer),
            сохраняя порядок абзацев и таблиц.
            """
            root = getattr(container._element, "body", None)
            root = root if root is not None else container._element

            for child in root.iterchildren():
                tag = child.tag
                if tag.endswith("}p"):
                    yield Paragraph(child, container)
                elif tag.endswith("}tbl"):
                    yield Table(child, container)

        def collect_text_from_table(table, mapping: List[Tuple[str, str]]) -> str:
            lines = []
            for row in table.rows:
                cells_text = []
                for cell in row.cells:
                    cell_pars = [
                        apply_mapping_to_text(p.text, mapping) for p in cell.paragraphs
                    ]
                    cells_text.append("\n".join(cell_pars).strip())
                lines.append("\t".join(cells_text))
            return "\n".join(lines)

        def collect_text_with_replacements(
            doc_obj, mapping: List[Tuple[str, str]]
        ) -> str:
            chunks: List[str] = []

            # Основное тело документа
            for block in iter_block_items(doc_obj):
                if isinstance(block, Paragraph):
                    chunks.append(apply_mapping_to_text(block.text, mapping))
                elif isinstance(block, Table):
                    chunks.append(collect_text_from_table(block, mapping))

            # Колонтитулы всех секций
            for section in doc_obj.sections:
                header_chunks: List[str] = []
                for block in iter_block_items(section.header):
                    if isinstance(block, Paragraph):
                        header_chunks.append(apply_mapping_to_text(block.text, mapping))
                    elif isinstance(block, Table):
                        header_chunks.append(collect_text_from_table(block, mapping))
                if header_chunks:
                    chunks.append("\n".join(header_chunks))

                footer_chunks: List[str] = []
                for block in iter_block_items(section.footer):
                    if isinstance(block, Paragraph):
                        footer_chunks.append(apply_mapping_to_text(block.text, mapping))
                    elif isinstance(block, Table):
                        footer_chunks.append(collect_text_from_table(block, mapping))
                if footer_chunks:
                    chunks.append("\n".join(footer_chunks))

            return "\n\n".join(chunks)

        try:
            doc = Document(str(file_path))
        except Exception as e:
            raise ValueError(f"Ошибка чтения .docx файла: {e}")

        # тут replacements уже вида ("[ЗАМЕНА001]", "Иванов")
        mapping = replacements
        text = collect_text_with_replacements(doc, mapping)
        return text

    def process_txt_file_deanonymize(
        self, file_path: Path, replacements: List[Tuple[str, str]]
    ) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="cp1251") as f:
                    content = f.read()
            except UnicodeDecodeError as e:
                raise ValueError(f"Не удалось прочитать файл: {e}")

        for pattern, original in replacements:
            content = content.replace(pattern, original)

        return content

    def process_docx_file_deanonymize(
        self, file_path: Path, replacements: List[Tuple[str, str]]
    ) -> str:
        try:
            return self.process_docx_file(file_path, replacements, anonymize=False)
        except Exception as e:
            raise ValueError(f"Ошибка деанонимизации .docx файла: {e}")

    def process_doc_file_deanonymize(
        self, file_path: Path, replacements: List[Tuple[str, str]]
    ) -> str:
        if textract is None:
            raise ValueError(
                "Поддержка .doc требует textract; "
                "конвертируйте в .docx или установите textract"
            )

        try:
            content = textract.process(str(file_path)).decode("utf-8")
        except Exception as e:
            raise ValueError(f"Ошибка чтения .doc файла: {e}")

        for pattern, original in replacements:
            content = content.replace(pattern, original)

        return content


class AnonymizationService:
    """
    Сервис анонимизации/деанонимизации.

    Использует:
    - PseudonymStore для получения списков замен;
    - FileProcessor для работы с файлами.

    Не знает ни про tkinter, ни про виджеты.
    """

    def __init__(self, store: PseudonymStore, file_processor: FileProcessor) -> None:
        self.store = store
        self.file_processor = file_processor

    def anonymize_file(self, filename: str) -> str:
        file_path = Path(filename)
        replacements = self.store.get_replacement_order()

        suffix = file_path.suffix.lower()
        if suffix == ".txt":
            return self.file_processor.process_txt_file(file_path, replacements)
        elif suffix == ".docx":
            return self.file_processor.process_docx_file(
                file_path, replacements, anonymize=True
            )
        elif suffix == ".doc":
            return self.file_processor.process_doc_file(
                file_path, replacements, anonymize=True
            )
        else:
            raise ValueError("Неподдерживаемый формат файла")

    def deanonymize_file(
        self,
        filename: str,
        *,
        to_clipboard: bool = True,
        to_file: bool = False,
    ) -> str:
        file_path = Path(filename)
        suffix = file_path.suffix.lower()
        replacements = self.store.get_deanonymization_order()

        # ---- режим "в файл" ----
        if to_file:
            if suffix == ".txt":
                content = self.file_processor.process_txt_file_deanonymize(
                    file_path, replacements
                )
                output_path = self.file_processor.get_output_path(file_path, "_деанон")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)
                return f"Файл сохранён как: {output_path.name}"

            if suffix == ".docx":
                # используем "анонимизацию наоборот":
                return self.file_processor._process_docx_with_mapping_and_save(
                    file_path=file_path,
                    mapping=replacements,
                    suffix="_деанон",
                )

            if suffix == ".doc":
                content = self.file_processor.process_doc_file_deanonymize(
                    file_path, replacements
                )
                output_path = self.file_processor.get_output_path(
                    file_path, "_деанон", ".txt"
                )
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)
                return (
                    f"Файл .doc деанонимизирован и "
                    f"сохранён как текст: {output_path.name}"
                )

            raise ValueError("Неподдерживаемый формат файла")

        # ---- режим "в буфер" / просто вернуть текст ----
        if suffix == ".txt":
            content = self.file_processor.process_txt_file_deanonymize(
                file_path, replacements
            )
        elif suffix == ".docx":
            content = self.file_processor.process_docx_file(
                file_path, replacements, anonymize=False
            )
        elif suffix == ".doc":
            content = self.file_processor.process_doc_file_deanonymize(
                file_path, replacements
            )
        else:
            raise ValueError("Неподдерживаемый формат файла")

        if to_clipboard:
            pyperclip.copy(content)
            return "Деанонимизированный текст скопирован в буфер обмена."

        return content


class AnonymizerApp:
    def __init__(self, root) -> None:
        self.entry_anon: Optional[ttk.Entry] = None
        self.entry_pseudonym: Optional[ttk.Entry] = None
        self.btn_add_pair: Optional[ttk.Button] = None
        self.frame1: Optional[LabelFrameTopLeft] = None
        self.frame2: Optional[LabelFrameBottomLeft] = None
        self.frame3: Optional[LabelFrameRight] = None
        self.config_file: Path = self.get_config_path()
        self.config = configparser.ConfigParser()
        self.root = root
        self.root.title("Анонимизатор текстов")
        self.root.geometry("640x480")

        self.setup_logging()
        self.load_config()

        self.store = PseudonymStore(self.get_pseudos_path())
        self.store.load()

        self.anon_dict = self.store.anon_dict
        self.deanon_dict = self.store.deanon_dict

        self.file_processor = FileProcessor()
        self.service = AnonymizationService(self.store, self.file_processor)

        self.setup_gui()
        self.center_or_restore_window()
        self.bind_events()

    @staticmethod
    def setup_logging():
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s - %(levelname)s - %(message)s",
            handlers=[
                logging.FileHandler("anonymizer.log", encoding="utf-8"),
                logging.StreamHandler(),
            ],
        )

    def load_config(self):
        if not self.config_file.exists():
            self.create_default_config()

        self.config.read(self.config_file, encoding="utf-8")

    @staticmethod
    def get_config_path() -> Path:
        if settings.main_ini_dir:
            return Path(settings.main_ini_dir) / "main.ini"
        return Path(__file__).parent / "main.ini"

    @staticmethod
    def get_pseudos_path() -> Path:
        if settings.pseudos_list_dir:
            return Path(settings.pseudos_list_dir) / "pseudos.json"
        return Path(__file__).parent / "pseudos.json"

    def create_default_config(self):
        self.config["window"] = {"coord_x": "100", "coord_y": "100"}
        self.config["config"] = {
            "load_dir": str(Path.cwd()),
            "last_open_dir": str(Path.cwd()),
        }
        with open(self.config_file, "w", encoding="utf-8") as f:
            self.config.write(f)

    def load_pseudos(self):
        self.store.load()

    def save_pseudos(self):
        self.store.save()

    def setup_gui(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(0, weight=1)

        left_frame = ttk.Frame(main_frame)
        left_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        left_frame.columnconfigure(0, weight=1)
        left_frame.rowconfigure(1, weight=1)

        self.setup_frame1(left_frame)
        self.setup_frame2(left_frame)
        self.setup_frame3(main_frame)

    def setup_frame1(self, parent):
        self.frame1 = LabelFrameTopLeft(parent, add_pair_func=self.add_pair)
        logging.info(self.frame1.__dict__)

    def setup_frame2(self, parent):
        self.frame2 = LabelFrameBottomLeft(
            parent,
            anonymize_func=self.anonymize_file,
            deanonymize_func=self.deanonymize_file,
            deanonymize_to_file_func=self.deanonymize_file_to_file,
            deanonymize_from_clipboard_func=self.deanonymize_clipboard,
        )

    def setup_frame3(self, parent):
        self.frame3 = LabelFrameRight(parent, self.delete_pair)
        self.update_anon_list()

    def bind_events(self):
        self.frame1.entry_anon.bind("<KeyRelease>", self.on_anon_entry_change)
        self.frame3.listbox_anon.bind("<<ListboxSelect>>", self.on_list_select)
        self.root.bind("<Configure>", self.on_window_move)

    def center_or_restore_window(self):
        try:
            x = self.config.getint("window", "coord_x")
            y = self.config.getint("window", "coord_y")
            self.root.geometry(f"+{x}+{y}")
        except (configparser.NoOptionError, ValueError):
            self.root.eval("tk::PlaceWindow . center")

    def on_window_move(self, event):
        if event.widget == self.root:
            self.config.set("window", "coord_x", str(self.root.winfo_x()))
            self.config.set("window", "coord_y", str(self.root.winfo_y()))
            with open(self.config_file, "w", encoding="utf-8") as f:
                self.config.write(f)

    def on_anon_entry_change(self, event):
        if self.frame1.entry_anon.get().strip():
            self.frame1.btn_add_pair.config(state="normal")
        else:
            self.frame1.btn_add_pair.config(state="disabled")

    def on_list_select(self, event):
        selection = self.frame3.listbox_anon.curselection()
        if selection:
            key = self.frame3.listbox_anon.get(selection[0])
            pseudonym = self.anon_dict.get(key, "")
            self.frame3.deanon_val_entry.config(state="normal")
            self.frame3.deanon_val_entry.delete(0, tk.END)
            self.frame3.deanon_val_entry.insert(0, pseudonym)
            self.frame3.deanon_val_entry.config(state="readonly")

    def generate_pseudonym(self, initial_pseudo: Optional[str] = None) -> str:
        return self.store.generate_pseudonym(initial_pseudo)

    def add_pair(self):
        key = self.frame1.entry_anon.get().strip()
        value = self.frame1.entry_pseudonym.get().strip()

        if not key:
            messagebox.showwarning("Предупреждение", "Введите текст для анонимизации")
            return

        if key in self.anon_dict:
            messagebox.showwarning("Предупреждение", "Такой ключ уже существует")
            return

        if not value:
            value = self.generate_pseudonym()

        if value in self.deanon_dict:
            value = self.generate_pseudonym(value)

        self.anon_dict[key] = value
        self.deanon_dict[value] = key

        self.save_pseudos()
        self.update_anon_list()

        self.frame1.entry_anon.delete(0, tk.END)
        self.frame1.entry_pseudonym.delete(0, tk.END)
        self.frame1.entry_anon.focus()

        for i, item in enumerate(self.frame3.listbox_anon.get(0, tk.END)):
            if item == key:
                self.frame3.listbox_anon.selection_clear(0, tk.END)
                self.frame3.listbox_anon.selection_set(i)
                self.frame3.listbox_anon.see(i)
                break

    def delete_pair(self):
        selection = self.frame3.listbox_anon.curselection()
        if not selection:
            return

        key = self.frame3.listbox_anon.get(selection[0])

        if not messagebox.askyesno(
            "Подтверждение", f"Удалить пару '{key}' -> '{self.anon_dict[key]}'?"
        ):
            return

        value = self.anon_dict[key]
        del self.anon_dict[key]
        del self.deanon_dict[value]

        self.save_pseudos()
        self.update_anon_list()
        self.frame1.entry_anon.focus()

    def update_anon_list(self):
        self.frame3.listbox_anon.delete(0, tk.END)
        for key in sorted(self.anon_dict.keys()):
            self.frame3.listbox_anon.insert(tk.END, key)

    def get_replacement_order(self) -> List[Tuple[str, str]]:
        return self.store.get_replacement_order()

    def get_deanonymization_order(self) -> List[Tuple[str, str]]:
        return self.store.get_deanonymization_order()

    def load_filename(self) -> str:
        load_dir = self.config.get("config", "last_open_dir", fallback=str(Path.cwd()))
        filename = filedialog.askopenfilename(
            initialdir=load_dir,
            filetypes=[
                ("Поддерживаемые файлы", ["*.txt", "*.doc", "*.docx"]),
                ("Все файлы", "*.*"),
            ],
        )
        self.config.set("config", "last_open_dir", str(Path(filename).parent))
        with open(self.config_file, "w", encoding="utf-8") as f:
            self.config.write(f)
        return filename

    def anonymize_file(self):
        filename = self.load_filename()

        if not filename:
            return

        def worker():
            try:
                self.root.config(cursor="watch")
                for btn in [
                    self.frame2.btn_anon,
                    self.frame2.btn_deanon_file,
                    self.frame2.btn_deanon_to_file,
                    self.frame2.btn_deanon_clipboard,
                ]:
                    btn.config(state="disabled")

                result = self.service.anonymize_file(filename)
                self.root.after(0, lambda: self.show_anonymization_result(result))
            except Exception as e:
                logging.error(f"Ошибка анонимизации: {str(e)}")
                self.root.after(0, lambda: messagebox.showerror("Ошибка", "str(e)"))
            finally:
                self.root.after(0, self.restore_ui_state)

        threading.Thread(target=worker, daemon=True).start()

    def show_anonymization_result(self, message: str):
        messagebox.showinfo("Успех", message)

    def restore_ui_state(self):
        self.root.config(cursor="")
        for btn in [
            self.frame2.btn_anon,
            self.frame2.btn_deanon_file,
            self.frame2.btn_deanon_to_file,
            self.frame2.btn_deanon_clipboard,
        ]:
            btn.config(state="normal")

    def deanonymize_file(self):
        filename = self.load_filename()

        if not filename:
            return

        def worker():
            try:
                self.root.config(cursor="watch")
                for btn in [
                    self.frame2.btn_anon,
                    self.frame2.btn_deanon_file,
                    self.frame2.btn_deanon_to_file,
                    self.frame2.btn_deanon_clipboard,
                ]:
                    btn.config(state="disabled")

                result = self.service.deanonymize_file(filename)
                self.root.after(0, lambda: self.show_deanonymization_result(result))
            except Exception as e:
                # self.root.after(0, lambda: messagebox.showerror("Ошибка", str(e)))
                self.root.after(0, messagebox.showerror, "Ошибка", str(e))
                raise e
            finally:
                self.root.after(0, self.restore_ui_state)

        threading.Thread(target=worker, daemon=True).start()

    def deanonymize_file_to_file(self):
        filename = self.load_filename()

        if not filename:
            return

        def worker():
            try:
                self.root.config(cursor="watch")
                for btn in [
                    self.frame2.btn_anon,
                    self.frame2.btn_deanon_file,
                    self.frame2.btn_deanon_to_file,
                    self.frame2.btn_deanon_clipboard,
                ]:
                    btn.config(state="disabled")

                result = self.service.deanonymize_file(
                    filename,
                    to_clipboard=False,
                    to_file=True,
                )
                self.root.after(0, lambda: self.show_deanonymization_result(result))
            except Exception as e:
                self.root.after(0, messagebox.showerror, "Ошибка", str(e))
                raise e
            finally:
                self.root.after(0, self.restore_ui_state)

        threading.Thread(target=worker, daemon=True).start()

    def show_deanonymization_result(self, message: str):
        messagebox.showinfo("Успех", message)

    def deanonymize_clipboard(self):
        try:
            clipboard_content = pyperclip.paste()
            if not clipboard_content.strip():
                messagebox.showwarning("Предупреждение", "Буфер обмена пуст")
                return
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать буфер обмена: {e}")
            return

        def worker():
            try:
                self.root.config(cursor="watch")
                for btn in [
                    self.frame2.btn_anon,
                    self.frame2.btn_deanon_file,
                    self.frame2.btn_deanon_to_file,
                    self.frame2.btn_deanon_clipboard,
                ]:
                    btn.config(state="disabled")

                replacements = self.get_deanonymization_order()
                content = clipboard_content

                for pattern, original in replacements:
                    content = content.replace(pattern, original)

                pyperclip.copy(content)
                self.root.after(
                    0,
                    lambda: messagebox.showinfo(
                        "Успех", "Деанонимизированный текст скопирован в буфер обмена."
                    ),
                )
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Ошибка", str(e)))
            finally:
                self.root.after(0, self.restore_ui_state)

        threading.Thread(target=worker, daemon=True).start()


def main():
    root = tk.Tk()
    app = AnonymizerApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
