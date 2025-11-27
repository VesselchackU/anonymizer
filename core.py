import json
import logging
import os
import shutil
import tempfile
from pathlib import Path
from tkinter import messagebox
from typing import Dict, List, Optional, Tuple

import pyperclip
import textract
from docx import Document


class PseudonymStore:
    """
    Класс для работы со словарём псевдонимов.

    Хранит:
    - путь к JSON-файлу словаря;
    - пары «оригинал → псевдоним» и «псевдоним → оригинал».

    Предоставляет методы для:
    - загрузки и сохранения словаря из/в файл;
    - генерации новых псевдонимов;
    - формирования упорядоченных списков замен для анонимизации и деанонимизации.
    """

    def __init__(self, path: Path) -> None:
        self.path = path
        self.anon_dict: Dict[str, str] = {}
        self.deanon_dict: Dict[str, str] = {}

    def load(self) -> None:
        pseudos_file = self.path
        try:
            if pseudos_file.exists():
                with open(pseudos_file, "r", encoding="utf-8") as f:
                    self.anon_dict = json.load(f)
                self.deanon_dict = {v: k for k, v in self.anon_dict.items()}
                logging.info(f"Загружено {len(self.anon_dict)} псевдонимов")
        except Exception as e:
            logging.error(f"Ошибка загрузки pseudos.json: {e}")
            messagebox.showerror(
                "Ошибка", f"Не удалось загрузить словарь псевдонимов: {e}"
            )

    def save(self) -> None:
        pseudos_file = self.path
        backup_file = pseudos_file.with_suffix(".json.bak")

        try:
            if pseudos_file.exists():
                shutil.copy2(pseudos_file, backup_file)

            with tempfile.NamedTemporaryFile(
                mode="w", encoding="utf-8", dir=pseudos_file.parent, delete=False
            ) as tmp:
                json.dump(self.anon_dict, tmp, ensure_ascii=False, indent=2)
                tmp_path = tmp.name

            os.replace(tmp_path, pseudos_file)
            logging.info("Словарь псевдонимов сохранён")
        except Exception as e:
            logging.error(f"Ошибка сохранения pseudos.json: {e}")
            messagebox.showerror(
                "Ошибка", f"Не удалось сохранить словарь псевдонимов: {e}"
            )

    def generate_pseudonym(self, initial_pseudo: Optional[str] = None) -> str:
        max_num = 0
        initial_pseudo = initial_pseudo or "ЗАМЕНА"
        for pseudo in self.deanon_dict.keys():
            if pseudo.startswith(initial_pseudo):
                try:
                    num_str = pseudo[len(initial_pseudo) :]
                    num = int(num_str)
                    max_num = max(max_num, num)
                except ValueError:
                    continue
        new_num = max_num + 1
        return f"{initial_pseudo}{new_num:03d}"

    def get_replacement_order(self) -> List[Tuple[str, str]]:
        items = list(self.anon_dict.items())
        items.sort(key=lambda x: (-len(x[0]), x[0]))
        return items

    def get_deanonymization_order(self) -> List[Tuple[str, str]]:
        items = [
            (f"[{pseudo}]", original) for pseudo, original in self.deanon_dict.items()
        ]
        items.sort(key=lambda x: (-len(x[0]), x[0]))
        return items


class FileProcessor:
    """
    Класс для работы с файлами (.txt, .docx, .doc).

    Отвечает за:
    - чтение исходного файла;
    - применение списка замен;
    - сохранение результата и генерацию имени выходного файла.
    """

    def __init__(self) -> None:
        # Пока без состояния, оставляем на будущее.
        pass

    def _process_docx_with_mapping_and_save(
        self,
        file_path: Path,
        mapping: List[Tuple[str, str]],
        suffix: str,
    ) -> str:
        def apply_mapping_to_text(text: str, mapping: List[Tuple[str, str]]) -> str:
            # mapping уже должен быть отсортирован снаружи
            for src, dst in mapping:
                if src:
                    text = text.replace(src, dst)
            return text

        def replace_in_paragraph(paragraph, mapping: List[Tuple[str, str]]):
            old = paragraph.text
            new = apply_mapping_to_text(old, mapping)
            if new != old:
                # убиваем старые run'ы, чтобы не ломать форматирование сильнее, чем надо
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = new
                else:
                    paragraph.add_run(new)

        def replace_in_table(table, mapping: List[Tuple[str, str]]):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p, mapping)

        try:
            doc = Document(str(file_path))
        except Exception as e:
            user_msg = (
                "Не удалось открыть файл как документ Word.\n\n"
                "Возможные причины:\n"
                "• файл открыт и ещё не полностью сохранён в другой программе;\n"
                "• файл удалён, переименован или перемещён;\n"
                "• файл повреждён и не читается как .docx.\n\n"
                "Проверьте, что файл существует, закрыт во всех программах "
                "и попробуйте ещё раз."
            )
            raise ValueError(user_msg)

        # Основной текст
        for p in doc.paragraphs:
            replace_in_paragraph(p, mapping)
        # Таблицы
        for table in doc.tables:
            replace_in_table(table, mapping)
        # Колонтитулы всех секций
        for section in doc.sections:
            for p in section.header.paragraphs:
                replace_in_paragraph(p, mapping)
            for p in section.footer.paragraphs:
                replace_in_paragraph(p, mapping)
            for table in section.header.tables:
                replace_in_table(table, mapping)
            for table in section.footer.tables:
                replace_in_table(table, mapping)

        output_path = self.get_output_path(file_path, suffix, ".docx")
        try:
            doc.save(str(output_path))
        except Exception as e:
            raise ValueError(f"Ошибка сохранения .docx файла: {e}")

        return f"Файл сохранён как: {output_path.name}"

    def process_doc_file(
        self,
        file_path: Path,
        replacements: List[Tuple[str, str]],
        anonymize: bool = True,
    ) -> str:
        if textract is None:
            raise ValueError(
                "Поддержка .doc требует textract; "
                "конвертируйте в .docx или установите textract"
            )

        try:
            text = textract.process(str(file_path)).decode("utf-8")
        except Exception as e:
            raise ValueError(f"Ошибка чтения .doc файла: {e}")

        if anonymize:
            for original, pseudo in replacements:
                text = text.replace(original, f"[{pseudo}]")
        else:
            for pseudo_pattern, original in replacements:
                text = text.replace(pseudo_pattern, original)

        suffix = "_аноним" if anonymize else "_деаном"
        output_path = self.get_output_path(file_path, suffix, ".txt")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

        return f"Файл сохранён как: {output_path.name}"

    def get_output_path(
        self, original_path: Path, suffix: str, force_ext: Optional[str] = None
    ) -> Path:
        ext = force_ext if force_ext else original_path.suffix
        base_name = original_path.stem

        counter = 0
        while True:
            if counter == 0:
                new_name = f"{base_name}{suffix}{ext}"
            else:
                new_name = f"{base_name}{suffix}{counter:03d}{ext}"

            output_path = original_path.parent / new_name
            if not output_path.exists():
                return output_path
            counter += 1

    def process_txt_file(
        self, file_path: Path, replacements: List[Tuple[str, str]]
    ) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="cp1251") as f:
                    content = f.read()
            except UnicodeDecodeError as e:
                raise ValueError(f"Не удалось прочитать файл: {e}")

        for original, pseudo in replacements:
            content = content.replace(original, f"[{pseudo}]")

        output_path = self.get_output_path(file_path, "_аноним")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

        return f"Файл сохранён как: {output_path.name}"

    def process_docx_file(
        self,
        file_path: Path,
        replacements: List[Tuple[str, str]],
        anonymize: bool = True,
    ) -> str:
        # anonymize = True  -> анонимизируем и сохраняем .docx
        # anonymize = False -> деанонимизируем и возвращаем текст

        if anonymize:
            # пары (оригинал, псевдоним) -> (оригинал, [псевдоним])
            mapping = [(orig, f"[{pseudo}]") for orig, pseudo in replacements]
            return self._process_docx_with_mapping_and_save(
                file_path=file_path,
                mapping=mapping,
                suffix="_аноним",
            )

        # -------- режим деанонимизации: вернуть текст в буфер --------
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        def apply_mapping_to_text(text: str, mapping: List[Tuple[str, str]]) -> str:
            for src, dst in mapping:
                if src:
                    text = text.replace(src, dst)
            return text

        def iter_block_items(container):
            """
            Идём по детям контейнера (документ, ячейка, header/footer),
            сохраняя порядок абзацев и таблиц.
            """
            root = getattr(container._element, "body", None)
            root = root if root is not None else container._element

            for child in root.iterchildren():
                tag = child.tag
                if tag.endswith("}p"):
                    yield Paragraph(child, container)
                elif tag.endswith("}tbl"):
                    yield Table(child, container)

        def collect_text_from_table(table, mapping: List[Tuple[str, str]]) -> str:
            lines = []
            for row in table.rows:
                cells_text = []
                for cell in row.cells:
                    cell_pars = [
                        apply_mapping_to_text(p.text, mapping) for p in cell.paragraphs
                    ]
                    cells_text.append("\n".join(cell_pars).strip())
                lines.append("\t".join(cells_text))
            return "\n".join(lines)

        def collect_text_with_replacements(
            doc_obj, mapping: List[Tuple[str, str]]
        ) -> str:
            chunks: List[str] = []

            # Основное тело документа
            for block in iter_block_items(doc_obj):
                if isinstance(block, Paragraph):
                    chunks.append(apply_mapping_to_text(block.text, mapping))
                elif isinstance(block, Table):
                    chunks.append(collect_text_from_table(block, mapping))

            # Колонтитулы всех секций
            for section in doc_obj.sections:
                header_chunks: List[str] = []
                for block in iter_block_items(section.header):
                    if isinstance(block, Paragraph):
                        header_chunks.append(apply_mapping_to_text(block.text, mapping))
                    elif isinstance(block, Table):
                        header_chunks.append(collect_text_from_table(block, mapping))
                if header_chunks:
                    chunks.append("\n".join(header_chunks))

                footer_chunks: List[str] = []
                for block in iter_block_items(section.footer):
                    if isinstance(block, Paragraph):
                        footer_chunks.append(apply_mapping_to_text(block.text, mapping))
                    elif isinstance(block, Table):
                        footer_chunks.append(collect_text_from_table(block, mapping))
                if footer_chunks:
                    chunks.append("\n".join(footer_chunks))

            return "\n\n".join(chunks)

        try:
            doc = Document(str(file_path))
        except Exception as e:
            raise ValueError(f"Ошибка чтения .docx файла: {e}")

        # тут replacements уже вида ("[ЗАМЕНА001]", "Иванов")
        mapping = replacements
        text = collect_text_with_replacements(doc, mapping)
        return text

    def process_txt_file_deanonymize(
        self, file_path: Path, replacements: List[Tuple[str, str]]
    ) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="cp1251") as f:
                    content = f.read()
            except UnicodeDecodeError as e:
                raise ValueError(f"Не удалось прочитать файл: {e}")

        for pattern, original in replacements:
            content = content.replace(pattern, original)

        return content

    def process_docx_file_deanonymize(
        self, file_path: Path, replacements: List[Tuple[str, str]]
    ) -> str:
        try:
            return self.process_docx_file(file_path, replacements, anonymize=False)
        except Exception as e:
            raise ValueError(f"Ошибка деанонимизации .docx файла: {e}")

    def process_doc_file_deanonymize(
        self, file_path: Path, replacements: List[Tuple[str, str]]
    ) -> str:
        if textract is None:
            raise ValueError(
                "Поддержка .doc требует textract; "
                "конвертируйте в .docx или установите textract"
            )

        try:
            content = textract.process(str(file_path)).decode("utf-8")
        except Exception as e:
            raise ValueError(f"Ошибка чтения .doc файла: {e}")

        for pattern, original in replacements:
            content = content.replace(pattern, original)

        return content


class AnonymizationService:
    """
    Сервис анонимизации/деанонимизации.

    Использует:
    - PseudonymStore для получения списков замен;
    - FileProcessor для работы с файлами.

    Не знает ни про tkinter, ни про виджеты.
    """

    def __init__(self, store: PseudonymStore, file_processor: FileProcessor) -> None:
        self.store = store
        self.file_processor = file_processor

    def anonymize_file(self, filename: str) -> str:
        file_path = Path(filename)
        replacements = self.store.get_replacement_order()

        suffix = file_path.suffix.lower()
        if suffix == ".txt":
            return self.file_processor.process_txt_file(file_path, replacements)
        elif suffix == ".docx":
            return self.file_processor.process_docx_file(
                file_path, replacements, anonymize=True
            )
        elif suffix == ".doc":
            return self.file_processor.process_doc_file(
                file_path, replacements, anonymize=True
            )
        else:
            raise ValueError("Неподдерживаемый формат файла")

    def deanonymize_file(
        self,
        filename: str,
        *,
        to_clipboard: bool = True,
        to_file: bool = False,
    ) -> str:
        file_path = Path(filename)
        suffix = file_path.suffix.lower()
        replacements = self.store.get_deanonymization_order()

        # ---- режим "в файл" ----
        if to_file:
            if suffix == ".txt":
                content = self.file_processor.process_txt_file_deanonymize(
                    file_path, replacements
                )
                output_path = self.file_processor.get_output_path(file_path, "_деанон")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)
                return f"Файл сохранён как: {output_path.name}"

            if suffix == ".docx":
                # используем "анонимизацию наоборот":
                return self.file_processor._process_docx_with_mapping_and_save(
                    file_path=file_path,
                    mapping=replacements,
                    suffix="_деанон",
                )

            if suffix == ".doc":
                content = self.file_processor.process_doc_file_deanonymize(
                    file_path, replacements
                )
                output_path = self.file_processor.get_output_path(
                    file_path, "_деанон", ".txt"
                )
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)
                return (
                    f"Файл .doc деанонимизирован и "
                    f"сохранён как текст: {output_path.name}"
                )

            raise ValueError("Неподдерживаемый формат файла")

        # ---- режим "в буфер" / просто вернуть текст ----
        if suffix == ".txt":
            content = self.file_processor.process_txt_file_deanonymize(
                file_path, replacements
            )
        elif suffix == ".docx":
            content = self.file_processor.process_docx_file(
                file_path, replacements, anonymize=False
            )
        elif suffix == ".doc":
            content = self.file_processor.process_doc_file_deanonymize(
                file_path, replacements
            )
        else:
            raise ValueError("Неподдерживаемый формат файла")

        if to_clipboard:
            pyperclip.copy(content)
            return "Деанонимизированный текст скопирован в буфер обмена."

        return content
